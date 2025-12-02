"""
Modified DOCX to HTML converter for A4 page format
This version outputs HTML that's optimized for A4 page splitting in the frontend
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html import escape


def docx_to_html(docx_path):
    """
    Convert DOCX file to HTML with proper structure for A4 page splitting.
    
    Returns HTML string with clean, block-level elements that can be easily split into pages.
    """
    doc = Document(docx_path)
    html_parts = []
    
    # Process each paragraph and table
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para = element
            html_parts.append(process_paragraph(para, doc))
        elif element.tag.endswith('tbl'):  # Table
            table = element
            html_parts.append(process_table(table, doc))
    
    # Join all HTML parts
    html_content = '\n'.join(html_parts)
    
    # Return clean HTML without body/head tags (frontend will add those)
    return html_content


def process_paragraph(para, doc):
    """Process a paragraph element and return HTML."""
    # Get paragraph properties
    para_obj = None
    for p in doc.paragraphs:
        if p._element == para:
            para_obj = p
            break
    
    if not para_obj:
        return ''
    
    # Check if paragraph is empty
    if not para_obj.text.strip():
        return '<p>&nbsp;</p>'
    
    # Get alignment
    align = para_obj.alignment
    align_style = ''
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        align_style = 'text-align: center;'
    elif align == WD_ALIGN_PARAGRAPH.RIGHT:
        align_style = 'text-align: right;'
    elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        align_style = 'text-align: justify;'
    
    # Get paragraph style (heading, etc.)
    style_name = para_obj.style.name.lower()
    tag = 'p'
    
    if 'heading 1' in style_name or 'heading1' in style_name:
        tag = 'h1'
    elif 'heading 2' in style_name or 'heading2' in style_name:
        tag = 'h2'
    elif 'heading 3' in style_name or 'heading3' in style_name:
        tag = 'h3'
    elif 'heading 4' in style_name or 'heading4' in style_name:
        tag = 'h4'
    elif 'heading 5' in style_name or 'heading5' in style_name:
        tag = 'h5'
    elif 'heading 6' in style_name or 'heading6' in style_name:
        tag = 'h6'
    elif 'list' in style_name:
        tag = 'li'
    
    # Process runs (text with formatting)
    inner_html = process_runs(para_obj.runs)
    
    # Build paragraph HTML
    style_attr = f' style="{align_style}"' if align_style else ''
    return f'<{tag}{style_attr}>{inner_html}</{tag}>'


def process_runs(runs):
    """Process paragraph runs and return HTML with formatting."""
    html_parts = []
    
    for run in runs:
        text = escape(run.text)
        if not text:
            continue
        
        # Build style string
        styles = []
        
        # Bold
        if run.bold:
            styles.append('font-weight: bold;')
        
        # Italic
        if run.italic:
            styles.append('font-style: italic;')
        
        # Underline
        if run.underline:
            styles.append('text-decoration: underline;')
        
        # Font size
        if run.font.size:
            size_pt = run.font.size.pt
            styles.append(f'font-size: {size_pt}pt;')
        
        # Font color
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            color = f'#{rgb:06x}'
            styles.append(f'color: {color};')
        
        # Build HTML
        if styles:
            style_str = ' '.join(styles)
            html_parts.append(f'<span style="{style_str}">{text}</span>')
        else:
            html_parts.append(text)
    
    return ''.join(html_parts)


def process_table(table, doc):
    """Process a table element and return HTML."""
    html_parts = ['<table style="width: 100%; border-collapse: collapse; margin: 10px 0;">']
    
    # Find the table object
    table_obj = None
    for t in doc.tables:
        if t._element == table:
            table_obj = t
            break
    
    if not table_obj:
        return ''
    
    # Process rows
    for row in table_obj.rows:
        html_parts.append('<tr>')
        
        for cell in row.cells:
            # Get cell content
            cell_text = cell.text.strip()
            cell_html = escape(cell_text) if cell_text else '&nbsp;'
            
            # Check if it's a header row (first row)
            is_header = row == table_obj.rows[0]
            tag = 'th' if is_header else 'td'
            
            html_parts.append(f'<{tag} style="border: 1px solid #ddd; padding: 8px;">{cell_html}</{tag}>')
        
        html_parts.append('</tr>')
    
    html_parts.append('</table>')
    return ''.join(html_parts)


def clean_html(html):
    """Clean up HTML - remove extra whitespace, fix common issues."""
    # Remove multiple consecutive newlines
    html = re.sub(r'\n{3,}', '\n\n', html)
    
    # Remove whitespace between tags
    html = re.sub(r'>\s+<', '><', html)
    
    # Ensure proper spacing around block elements
    html = re.sub(r'(</(?:p|div|h[1-6]|li|table)>)', r'\1\n', html)
    html = re.sub(r'(<(?:p|div|h[1-6]|li|table)[^>]*>)', r'\n\1', html)
    
    return html.strip()


# Example usage function (modify based on your API endpoint)
def convert_docx_to_html_api(docx_file_path):
    """
    Convert DOCX file to HTML for API response.
    This is the main function you should call from your API endpoint.
    """
    try:
        # Convert DOCX to HTML
        html_content = docx_to_html(docx_file_path)
        
        # Clean up HTML
        html_content = clean_html(html_content)
        
        # Return just the body content (frontend will wrap it)
        return html_content
        
    except Exception as e:
        # Log error and return error message
        print(f"Error converting DOCX to HTML: {e}")
        return f'<p style="color: red;">Error converting document: {escape(str(e))}</p>'


# If this is used as a Flask/FastAPI endpoint, modify accordingly:
"""
Example Flask endpoint:

@app.route('/api/preview/docx', methods=['POST'])
def preview_docx():
    if 'file' not in request.files:
        return 'No file provided', 400
    
    file = request.files['file']
    if file.filename == '':
        return 'No file selected', 400
    
    # Save temporarily
    temp_path = os.path.join('/tmp', file.filename)
    file.save(temp_path)
    
    try:
        # Convert to HTML
        html_content = convert_docx_to_html_api(temp_path)
        
        # Return HTML
        return html_content, 200, {'Content-Type': 'text/html; charset=utf-8'}
    finally:
        # Clean up
        if os.path.exists(temp_path):
            os.remove(temp_path)
"""


#!/usr/bin/env python3
"""
Optimized CSV to DOC Converter
High-performance CSV to DOC conversion with optimizations for large files.
"""

import pandas as pd
import argparse
import os
import sys
from datetime import datetime
import traceback
from io import StringIO
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import psutil

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import qn as qn_ns
except ImportError as e:
    print(f"ERROR: Required DOCX library not available: {e}")
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

def process_chunk_threaded(chunk_df, chunk_info):
    """
    Process a chunk of data using threading.
    Returns processed data ready for DOC insertion.
    """
    try:
        print(f"Thread processing chunk {chunk_info}: {len(chunk_df)} rows")
        
        # Convert chunk to list of lists for DOC table
        chunk_rows = []
        for _, row in chunk_df.iterrows():
            chunk_rows.append(row.tolist())
        
        print(f"Thread completed chunk {chunk_info}: {len(chunk_rows)} rows processed")
        return {
            'chunk_info': chunk_info,
            'rows': chunk_rows,
            'columns': chunk_df.columns.tolist()
        }
    except Exception as e:
        print(f"Error processing chunk {chunk_info}: {e}")
        return None

def create_doc_from_csv_optimized(csv_file, output_file, title="CSV Data", author="Unknown", chunk_size=500, use_multiprocessing=True):
    """
    Optimized CSV to DOC conversion with performance improvements and multiprocessing.
    
    Args:
        csv_file (str): Path to input CSV file
        output_file (str): Path to output DOC file
        title (str): Document title
        author (str): Document author
        chunk_size (int): Number of rows to process at once
        use_multiprocessing (bool): Whether to use multiprocessing for parallel processing
    """
    print(f"Starting optimized CSV to DOC conversion...")
    print(f"Input: {csv_file}")
    print(f"Output: {output_file}")
    print(f"Chunk size: {chunk_size}")
    
    try:
        # Read CSV file with optimizations
        print("Reading CSV file with optimizations...")
        
        # Get file size for progress tracking
        file_size = os.path.getsize(csv_file)
        print(f"File size: {file_size / (1024*1024):.2f} MB")
        
        # Read CSV with optimized settings for large files
        print("Reading large CSV file with memory optimizations...")
        df = pd.read_csv(
            csv_file,
            dtype=str,  # Read all as strings to avoid type inference overhead
            na_filter=False,  # Disable NaN filtering for speed
            low_memory=False,  # Use more memory for speed
            chunksize=None,  # Read entire file at once for processing
            engine='c'  # Use C engine for better performance
        )
        
        print(f"CSV loaded: {len(df)} rows, {len(df.columns)} columns")
        
        # Create DOCX document
        print("Creating optimized DOCX document...")
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()
        
        # Add minimal title section
        print("Adding document header...")
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add creation info
        info_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Rows: {len(df)} | Columns: {len(df.columns)}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line
        
        # Create table with optimizations
        print("Creating optimized data table...")
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Add header row with minimal styling
        print("Adding table headers...")
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            # Minimal header styling
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)  # Smaller font for speed
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process data in chunks for better performance
        print(f"Processing {len(df)} rows in chunks of {chunk_size}...")
        print(f"Threading enabled: {use_multiprocessing}")
        print(f"Total rows: {len(df)}")
        print(f"Estimated processing time: {len(df) // 1000} seconds for large file")
        
        total_rows = len(df)
        
        # Use threading for better CPU utilization
        if use_multiprocessing and total_rows > 100:  # Reasonable threshold for threading
            try:
                cpu_count = psutil.cpu_count(logical=True)
                max_workers = min(cpu_count * 3, 24)  # Use more threads for I/O bound tasks
                print(f"Using threading with {max_workers} workers (CPU cores: {cpu_count})")
                
                # Split data into chunks for parallel processing
                chunk_data = []
                for chunk_start in range(0, total_rows, chunk_size):
                    chunk_end = min(chunk_start + chunk_size, total_rows)
                    chunk_df = df.iloc[chunk_start:chunk_end]
                    chunk_data.append((chunk_df, chunk_start))
                
                # Process chunks in parallel using threads
                print(f"Starting threaded processing of {len(chunk_data)} chunks...")
                processed_chunks = []
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    print(f"ThreadPoolExecutor created with {max_workers} workers")
                    # Submit all chunks
                    future_to_chunk = {
                        executor.submit(process_chunk_threaded, chunk_df, i): i 
                        for i, (chunk_df, _) in enumerate(chunk_data)
                    }
                    
                    # Collect results as they complete
                    for future in as_completed(future_to_chunk):
                        chunk_idx = future_to_chunk[future]
                        try:
                            result = future.result()
                            if result:
                                processed_chunks.append((chunk_idx, result))
                                print(f"Completed chunk {chunk_idx}")
                        except Exception as e:
                            print(f"Chunk {chunk_idx} failed: {e}")
                
                # Sort chunks by index to maintain order
                processed_chunks.sort(key=lambda x: x[0])
                
                # Add processed data to table
                print("Adding processed data to document...")
                for chunk_idx, chunk_result in processed_chunks:
                    for row_data in chunk_result['rows']:
                        table.add_row()
                        row_cells = table.rows[-1].cells
                        
                        for i, value in enumerate(row_data):
                            cell_value = str(value) if value else ""
                            row_cells[i].text = cell_value
                            
                            # Minimal styling for first few rows
                            if chunk_idx < 2:  # Style first 2 chunks
                                for paragraph in row_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(8)
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            except Exception as e:
                print(f"Error with threading: {e}, falling back to single-threaded")
                use_multiprocessing = False
        
        # Single-threaded processing (for small datasets or when threading fails)
        if not (use_multiprocessing and total_rows > 100):
            print("Using single-threaded processing...")
            print(f"Reason: use_multiprocessing={use_multiprocessing}, total_rows={total_rows}, threshold=100")
            
            # Pre-allocate table rows for better performance
            if total_rows > 1000:
                print("Pre-allocating table rows for large dataset...")
                for _ in range(min(chunk_size, total_rows)):
                    table.add_row()
            
            # Process data in chunks
            for chunk_start in range(0, total_rows, chunk_size):
                chunk_end = min(chunk_start + chunk_size, total_rows)
                chunk_df = df.iloc[chunk_start:chunk_end]
                
                print(f"Processing rows {chunk_start + 1}-{chunk_end} of {total_rows}")
                
                for idx, (_, row) in enumerate(chunk_df.iterrows()):
                    # Add row if not pre-allocated
                    if total_rows <= 1000 or idx >= chunk_size:
                        table.add_row()
                    
                    row_cells = table.rows[-1].cells
                    
                    # Process cells with minimal styling
                    for i, value in enumerate(row):
                        # Convert to string efficiently
                        cell_value = str(value) if value else ""
                        row_cells[i].text = cell_value
                        
                        # Minimal cell styling (only for first few rows to save time)
                        if chunk_start < 100:  # Only style first 100 rows
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)  # Smaller font
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add minimal summary
        print("Adding document summary...")
        doc.add_paragraph()
        summary_para = doc.add_paragraph(f"Total: {len(df)} rows × {len(df.columns)} columns")
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Memory cleanup for large files
        print("Cleaning up memory...")
        del df  # Free up memory from the large DataFrame
        
        # Save document with optimizations
        print(f"Saving DOCX document...")
        doc.save(output_file)
        
        # Verify file was created
        if os.path.exists(output_file):
            file_size = os.path.getsize(output_file)
            print(f"DOCX file created successfully: {file_size / (1024*1024):.2f} MB")
            return True
        else:
            print("ERROR: DOCX file was not created")
            return False
            
    except Exception as e:
        print(f"ERROR: Failed to create DOCX from CSV: {e}")
        traceback.print_exc()
        return False

def main():
    parser = argparse.ArgumentParser(description='Convert CSV to DOC format (optimized)')
    parser.add_argument('csv_file', help='Input CSV file path')
    parser.add_argument('output_file', help='Output DOC file path')
    parser.add_argument('--title', default='CSV Data', help='Document title')
    parser.add_argument('--author', default='Unknown', help='Document author')
    parser.add_argument('--chunk-size', type=int, default=1000, help='Chunk size for processing large files')
    
    args = parser.parse_args()
    
    print("=== Optimized CSV to DOC Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    # Check if input file exists
    if not os.path.exists(args.csv_file):
        print(f"ERROR: Input CSV file not found: {args.csv_file}")
        sys.exit(1)
    
    # Check pandas availability
    try:
        print(f"Pandas version: {pd.__version__}")
    except Exception as e:
        print(f"ERROR: Pandas not available: {e}")
        sys.exit(1)
    
    # Check python-docx availability
    try:
        from docx import Document
        print("python-docx library is available")
    except Exception as e:
        print(f"ERROR: python-docx not available: {e}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(args.output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Convert CSV to DOCX with optimizations
    success = create_doc_from_csv_optimized(
        args.csv_file,
        args.output_file,
        args.title,
        args.author,
        args.chunk_size
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)

if __name__ == "__main__":
    main()

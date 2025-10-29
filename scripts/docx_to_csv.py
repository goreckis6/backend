#!/usr/bin/env python3
"""
DOCX to CSV Converter
Extracts text content and tables from DOCX files and converts to CSV format
"""

import os
import sys
import argparse
import csv
import traceback
from docx import Document
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import pandas as pd


def clean_text(text):
    """Clean and normalize text content"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = ' '.join(text.split())
    # Remove leading/trailing whitespace
    text = text.strip()
    return text


def extract_tables_from_docx(docx_file):
    """
    Extract all tables from DOCX file and return as list of DataFrames
    
    Args:
        docx_file (str): Path to input DOCX file
        
    Returns:
        list: List of pandas DataFrames, one for each table
    """
    try:
        doc = Document(docx_file)
        tables_data = []
        
        print(f"Found {len(doc.tables)} table(s) in document")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"Processing table {table_idx + 1}...")
            
            # Convert table to list of lists
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    row_data.append(cell_text)
                
                # Only add non-empty rows
                if any(cell for cell in row_data):
                    table_data.append(row_data)
            
            if table_data:
                # Create DataFrame
                # Use first row as headers if it exists
                if len(table_data) > 0:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if len(table_data) > 1 else None)
                    tables_data.append(df)
                    print(f"Table {table_idx + 1}: {len(df)} rows, {len(df.columns)} columns")
        
        return tables_data
        
    except Exception as e:
        print(f"ERROR: Failed to extract tables from DOCX: {e}")
        traceback.print_exc()
        return None


def extract_text_paragraphs_from_docx(docx_file):
    """
    Extract all paragraphs from DOCX file
    
    Args:
        docx_file (str): Path to input DOCX file
        
    Returns:
        list: List of paragraph texts
    """
    try:
        doc = Document(docx_file)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = clean_text(para.text)
            if text:
                paragraphs.append(text)
        
        print(f"Extracted {len(paragraphs)} paragraphs from document")
        return paragraphs
        
    except Exception as e:
        print(f"ERROR: Failed to extract paragraphs from DOCX: {e}")
        traceback.print_exc()
        return None


def convert_docx_to_csv(docx_file, output_file, extract_tables=True, include_paragraphs=True, delimiter=','):
    """
    Convert DOCX file to CSV format
    
    Args:
        docx_file (str): Path to input DOCX file
        output_file (str): Path to output CSV file
        extract_tables (bool): Extract tables from DOCX
        include_paragraphs (bool): Include paragraph text in CSV
        delimiter (str): CSV delimiter character
        
    Returns:
        bool: True if conversion successful, False otherwise
    """
    print(f"Starting DOCX to CSV conversion...")
    print(f"Input: {docx_file}")
    print(f"Output: {output_file}")
    print(f"Extract tables: {extract_tables}")
    print(f"Include paragraphs: {include_paragraphs}")
    print(f"Delimiter: {repr(delimiter)}")
    
    try:
        # Check if DOCX file exists
        if not os.path.exists(docx_file):
            print(f"ERROR: DOCX file does not exist: {docx_file}")
            return False
        
        file_size = os.path.getsize(docx_file)
        print(f"DOCX file size: {file_size} bytes")
        
        if file_size == 0:
            print("ERROR: Input file is empty")
            return False
        
        # Extract content
        all_rows = []
        
        # Extract tables if requested
        if extract_tables:
            tables = extract_tables_from_docx(docx_file)
            if tables:
                print(f"Found {len(tables)} table(s)")
                
                # Combine all tables into one CSV
                for table_idx, df in enumerate(tables):
                    if table_idx == 0:
                        # First table: use its structure
                        for _, row in df.iterrows():
                            all_rows.append(row.tolist())
                        # Add column names as first row
                        if len(all_rows) > 0:
                            all_rows.insert(0, df.columns.tolist())
                    else:
                        # Subsequent tables: append with separator
                        # Add empty row as separator
                        all_rows.append([])
                        # Add table identifier
                        all_rows.append([f"Table {table_idx + 1}"])
                        # Add column headers
                        all_rows.append(df.columns.tolist())
                        # Add data rows
                        for _, row in df.iterrows():
                            all_rows.append(row.tolist())
        
        # Extract paragraphs if requested and no tables found
        if include_paragraphs and (not extract_tables or not tables or len(tables) == 0):
            paragraphs = extract_text_paragraphs_from_docx(docx_file)
            if paragraphs:
                print(f"Using {len(paragraphs)} paragraphs as CSV rows")
                
                # Create CSV with paragraph number and content
                all_rows.append(['Paragraph Number', 'Content'])
                for idx, para in enumerate(paragraphs, 1):
                    all_rows.append([str(idx), para])
        
        # If no content was extracted
        if not all_rows:
            print("ERROR: No content extracted from DOCX file")
            print("Tried to extract tables:", extract_tables)
            print("Tried to extract paragraphs:", include_paragraphs)
            return False
        
        # Write to CSV file
        print(f"Writing CSV file with {len(all_rows)} rows...")
        
        with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile, delimiter=delimiter, quoting=csv.QUOTE_MINIMAL)
            
            for row in all_rows:
                writer.writerow(row)
        
        # Verify output file
        if os.path.exists(output_file):
            output_size = os.path.getsize(output_file)
            print(f"CSV file created successfully: {output_size} bytes")
            
            # Verify it's a valid CSV
            try:
                with open(output_file, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    row_count = sum(1 for row in reader)
                    print(f"Verified CSV file with {row_count} rows")
                    return True
            except Exception as verify_error:
                print(f"ERROR: Output file is not a valid CSV: {verify_error}")
                return False
        else:
            print("ERROR: CSV file was not created")
            return False
            
    except Exception as e:
        print(f"ERROR: Failed to convert DOCX to CSV: {e}")
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description='Convert DOCX file to CSV format')
    parser.add_argument('docx_file', help='Path to input DOCX file')
    parser.add_argument('output_file', help='Path to output CSV file')
    parser.add_argument('--no-tables', action='store_true',
                        help='Do not extract tables from DOCX')
    parser.add_argument('--no-paragraphs', action='store_true',
                        help='Do not include paragraphs in CSV')
    parser.add_argument('--delimiter', default=',',
                        help='CSV delimiter character (default: comma)')
    
    args = parser.parse_args()
    
    print("=== DOCX to CSV Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    success = convert_docx_to_csv(
        args.docx_file, 
        args.output_file,
        extract_tables=not args.no_tables,
        include_paragraphs=not args.no_paragraphs,
        delimiter=args.delimiter
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)


if __name__ == '__main__':
    main()


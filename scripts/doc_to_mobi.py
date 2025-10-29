#!/usr/bin/env python3
"""
DOC to MOBI Converter
Converts Microsoft Word DOC files to MOBI format for Kindle e-readers
Uses LibreOffice to convert DOC to DOCX, then python-docx + ebooklib for EPUB, and Calibre ebook-convert for MOBI
"""

import os
import sys
import argparse
import traceback
import subprocess
import tempfile
import shutil
from datetime import datetime
try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    HAS_DOCX_EPUB = True
except ImportError:
    HAS_DOCX_EPUB = False
import re


def clean_text(text):
    """Clean and normalize text content"""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text


def extract_text_from_paragraph(para):
    """Extract text from a paragraph, handling runs and formatting"""
    text_parts = []
    for run in para.runs:
        if run.text:
            text_parts.append(run.text)
    full_text = ''.join(text_parts)
    return clean_text(full_text)


def paragraph_to_html(para, heading_level=None):
    """Convert a paragraph to HTML"""
    text = extract_text_from_paragraph(para)
    if not text:
        return ""
    if heading_level:
        return f"<h{heading_level}>{text}</h{heading_level}>"
    style_name = para.style.name.lower() if para.style else ''
    if 'heading' in style_name or 'title' in style_name:
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            level = min(int(match.group(1)), 6)
            return f"<h{level}>{text}</h{level}>"
        else:
            return f"<h1>{text}</h1>"
    return f"<p>{text}</p>"


def table_to_html(table):
    """Convert a table to HTML"""
    html = "<table>"
    for row_idx, row in enumerate(table.rows):
        html += "<tr>"
        for cell in row.cells:
            cell_text = clean_text(cell.text)
            tag = "th" if row_idx == 0 else "td"
            html += f"<{tag}>{cell_text}</{tag}>"
        html += "</tr>"
    html += "</table>"
    return html


def extract_images_from_docx(docx_file, output_dir):
    """Extract images from DOCX file and save to output directory"""
    images = []
    doc = Document(docx_file)
    relationships = doc.part.rels
    for rel_id, rel in relationships.items():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_data = image_part.blob
                image_ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                image_filename = f"image_{len(images)}.{image_ext}"
                image_path = os.path.join(output_dir, image_filename)
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                images.append({
                    'filename': image_filename,
                    'path': image_path,
                    'mimetype': f"image/{image_ext}" if image_ext in ['png', 'jpg', 'jpeg', 'gif'] else "image/png"
                })
            except Exception as e:
                print(f"Warning: Could not extract image {rel_id}: {e}")
    return images


def find_libreoffice():
    """Find LibreOffice binary"""
    libreoffice_paths = [
        'libreoffice',
        '/usr/bin/libreoffice',
        '/usr/local/bin/libreoffice',
        '/opt/libreoffice/program/soffice',
        'soffice'
    ]
    for path in libreoffice_paths:
        try:
            result = subprocess.run(
                [path, '--version'],
                capture_output=True,
                check=True,
                timeout=5
            )
            print(f"Found LibreOffice at: {path}")
            return path
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


def convert_doc_to_docx_with_libreoffice(doc_file, output_dir):
    """Convert DOC to DOCX using LibreOffice"""
    libreoffice = find_libreoffice()
    if not libreoffice:
        return None
    
    base_name = os.path.splitext(os.path.basename(doc_file))[0]
    output_docx = os.path.join(output_dir, f"{base_name}.docx")
    
    cmd = [
        libreoffice,
        '--headless',
        '--invisible',
        '--nocrashreport',
        '--nodefault',
        '--nofirststartwizard',
        '--nolockcheck',
        '--nologo',
        '--norestore',
        '--convert-to', 'docx',
        '--outdir', output_dir,
        doc_file
    ]
    
    env = os.environ.copy()
    env['SAL_USE_VCLPLUGIN'] = 'svp'
    env['HOME'] = '/tmp'
    env['LANG'] = 'en_US.UTF-8'
    env['LC_ALL'] = 'en_US.UTF-8'
    
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=300,
            env=env,
            encoding='utf-8',
            errors='replace'
        )
        if os.path.exists(output_docx):
            return output_docx
        else:
            print(f"LibreOffice conversion failed: {result.stderr}")
            return None
    except Exception as e:
        print(f"Error converting DOC to DOCX: {e}")
        return None


def create_epub_from_docx(docx_file, epub_file, include_images=True, preserve_formatting=True, generate_toc=True):
    """Create EPUB file from DOCX file"""
    print(f"Creating EPUB from DOCX: {docx_file}")
    
    try:
        doc = Document(docx_file)
        book = epub.EpubBook()
        
        title = "Untitled Document"
        try:
            if doc.core_properties.title:
                title = doc.core_properties.title
        except:
            pass
        
        if not title or title == "":
            title = os.path.splitext(os.path.basename(docx_file))[0]
        
        book.set_identifier(os.path.basename(docx_file))
        book.set_title(title)
        
        author = "Unknown"
        try:
            if doc.core_properties.author:
                author = doc.core_properties.author
        except:
            pass
        book.add_author(author)
        book.set_language('en')
        
        image_dir = os.path.dirname(epub_file)
        images = []
        if include_images:
            images = extract_images_from_docx(docx_file, image_dir)
            for img in images:
                img_item = epub.EpubItem(uid=img['filename'], file_name=f"images/{img['filename']}", media_type=img['mimetype'], content=open(img['path'], 'rb').read())
                book.add_item(img_item)
        
        chapter_html_parts = []
        for block in doc.element.body:
            if isinstance(block, CT_P):
                paragraph = Paragraph(block, doc)
                chapter_html_parts.append(paragraph_to_html(paragraph))
            elif isinstance(block, CT_Tbl):
                table = Table(block, doc)
                chapter_html_parts.append(table_to_html(table))
        
        main_chapter_content = "".join(chapter_html_parts)
        
        if not main_chapter_content.strip():
            c1 = epub.EpubHtml(title='No Content', file_name='no_content.xhtml', lang='en')
            c1.content = '<h1>No Content</h1><p>The DOCX file did not contain extractable content.</p>'
            book.add_item(c1)
            book.toc = (epub.Link('no_content.xhtml', 'No Content', 'no_content'),)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ['nav', c1]
        else:
            c1 = epub.EpubHtml(title=title, file_name='chap_01.xhtml', lang='en')
            c1.content = main_chapter_content
            book.add_item(c1)
            
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            if generate_toc:
                soup = BeautifulSoup(main_chapter_content, 'html.parser')
                toc_items = []
                for i, heading in enumerate(soup.find_all(['h1', 'h2', 'h3']), 1):
                    heading_id = f"heading_{i}"
                    heading['id'] = heading_id
                    toc_items.append(epub.Link(f'chap_01.xhtml#{heading_id}', heading.get_text(strip=True), heading_id))
                book.toc = tuple(toc_items) if toc_items else (epub.Link('chap_01.xhtml', title, 'intro'),)
                c1.content = str(soup)
            else:
                book.toc = (epub.Link('chap_01.xhtml', title, 'intro'),)
            
            book.spine = ['nav', c1]
        
        epub.write_epub(epub_file, book, {})
        print(f"EPUB file created: {epub_file}")
        return True
    except Exception as e:
        print(f"Error creating EPUB: {e}")
        traceback.print_exc()
        return False


def find_ebook_convert():
    """Find Calibre ebook-convert binary"""
    ebook_convert_paths = [
        'ebook-convert',
        '/usr/bin/ebook-convert',
        '/usr/local/bin/ebook-convert',
        '/opt/calibre/bin/ebook-convert'
    ]
    for path in ebook_convert_paths:
        try:
            subprocess.run([path, '--version'], capture_output=True, check=True, timeout=10)
            print(f"Found ebook-convert at: {path}")
            return path
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


def convert_doc_to_mobi(doc_file, output_file, include_images=True, preserve_formatting=True, generate_toc=True, kindle_optimized=True):
    """
    Convert DOC file to MOBI format.
    Strategy: DOC -> DOCX (LibreOffice) -> EPUB (python-docx + ebooklib) -> MOBI (ebook-convert)
    """
    print(f"Starting DOC to MOBI conversion...")
    print(f"Input: {doc_file}")
    print(f"Output: {output_file}")
    print(f"Include images: {include_images}")
    print(f"Preserve formatting: {preserve_formatting}")
    print(f"Generate TOC: {generate_toc}")
    print(f"Kindle Optimized: {kindle_optimized}")

    if not os.path.exists(doc_file):
        print(f"ERROR: Input DOC file not found: {doc_file}")
        return False
    
    if os.path.getsize(doc_file) == 0:
        print(f"ERROR: Input DOC file is empty: {doc_file}")
        return False

    temp_dir = None
    temp_docx = None
    temp_epub = None
    
    try:
        temp_dir = tempfile.mkdtemp()
        print(f"Created temporary directory: {temp_dir}")

        # Step 1: Convert DOC to DOCX using LibreOffice
        print("Step 1: Converting DOC to DOCX using LibreOffice...")
        temp_docx = convert_doc_to_docx_with_libreoffice(doc_file, temp_dir)
        
        if not temp_docx or not os.path.exists(temp_docx):
            print("ERROR: Failed to convert DOC to DOCX using LibreOffice")
            return False
        
        print(f"Step 1 complete: DOCX created at {temp_docx}")
        
        # Step 2: Convert DOCX to EPUB using python-docx and ebooklib
        if not HAS_DOCX_EPUB:
            print("ERROR: Required libraries (python-docx, ebooklib, beautifulsoup4) are not available")
            return False
        
        print("Step 2: Converting DOCX to EPUB...")
        temp_epub = os.path.join(temp_dir, os.path.basename(temp_docx).replace('.docx', '.epub'))
        
        if not create_epub_from_docx(temp_docx, temp_epub, include_images, preserve_formatting, generate_toc):
            print("ERROR: Failed to convert DOCX to EPUB")
            return False
        
        if not os.path.exists(temp_epub) or os.path.getsize(temp_epub) == 0:
            print(f"ERROR: Intermediate EPUB file was not created or is empty: {temp_epub}")
            return False
        
        print(f"Step 2 complete: EPUB created at {temp_epub}")
        
        # Step 3: Convert EPUB to MOBI using ebook-convert (Calibre)
        print("Step 3: Converting EPUB to MOBI using ebook-convert...")
        ebook_convert = find_ebook_convert()
        
        if not ebook_convert:
            print("ERROR: ebook-convert (Calibre) is required for MOBI conversion but not available on PATH.")
            return False
        
        # Create output directory if needed
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        convert_cmd = [
            ebook_convert,
            temp_epub,
            output_file,
            '--output-profile', 'kindle' if kindle_optimized else 'default',
            '--mobi-file-type', 'new',  # 'new' for KF8, 'old' for Mobipocket
            '--no-default-epub-cover'  # Prevent Calibre from adding a generic cover
        ]
        
        # Extract title and author from DOCX for Calibre metadata
        try:
            doc = Document(temp_docx)
            title = doc.core_properties.title if doc.core_properties.title else os.path.basename(doc_file).replace('.doc', '')
            author = doc.core_properties.author if doc.core_properties.author else "Unknown"
            convert_cmd.extend(['--title', title, '--authors', author])
        except:
            pass
        
        print(f"Running command: {' '.join(convert_cmd)}")
        result = subprocess.run(convert_cmd, capture_output=True, text=True, timeout=300)
        
        if result.returncode != 0:
            print(f"ERROR: ebook-convert failed with return code {result.returncode}")
            print(f"stderr: {result.stderr}")
            raise RuntimeError(f"ebook-convert failed: {result.stderr}")
        
        print("MOBI conversion completed successfully")
        
        if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
            print(f"Successfully converted DOC to MOBI: {output_file}")
            return True
        else:
            print(f"ERROR: MOBI file was not created or is empty: {output_file}")
            return False

    except Exception as e:
        print(f"ERROR: Failed to convert DOC to MOBI: {e}")
        traceback.print_exc()
        return False
    finally:
        # Clean up temporary files
        if temp_epub and os.path.exists(temp_epub):
            try:
                os.unlink(temp_epub)
                print(f"Cleaned up temporary EPUB file: {temp_epub}")
            except Exception as e:
                print(f"Warning: Failed to clean up temporary EPUB file {temp_epub}: {e}")
        
        if temp_docx and os.path.exists(temp_docx):
            try:
                os.unlink(temp_docx)
                print(f"Cleaned up temporary DOCX file: {temp_docx}")
            except Exception as e:
                print(f"Warning: Failed to clean up temporary DOCX file {temp_docx}: {e}")
        
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                print(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as e:
                print(f"Warning: Failed to clean up temporary directory {temp_dir}: {e}")


def main():
    parser = argparse.ArgumentParser(description='Convert DOC file to MOBI format using Calibre ebook-convert')
    parser.add_argument('doc_file', help='Path to input DOC file')
    parser.add_argument('output_file', help='Path to output MOBI file')
    parser.add_argument('--no-images', action='store_true',
                        help='Exclude images from MOBI')
    parser.add_argument('--no-formatting', action='store_true',
                        help='Exclude formatting from MOBI')
    parser.add_argument('--no-toc', action='store_true',
                        help='Do not generate table of contents')
    parser.add_argument('--no-kindle-optimize', action='store_true',
                        help='Do not optimize MOBI for Kindle devices')
    
    args = parser.parse_args()
    
    print("=== DOC to MOBI Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    success = convert_doc_to_mobi(
        args.doc_file,
        args.output_file,
        include_images=not args.no_images,
        preserve_formatting=not args.no_formatting,
        generate_toc=not args.no_toc,
        kindle_optimized=not args.no_kindle_optimize
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)


if __name__ == '__main__':
    main()


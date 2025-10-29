#!/usr/bin/env python3
"""
DOCX to MOBI Converter
Converts Microsoft Word DOCX files to MOBI format for Kindle e-readers
Uses python-docx to extract content, ebooklib to create EPUB, and Calibre ebook-convert to create MOBI
"""

import os
import sys
import argparse
import traceback
import subprocess
import tempfile
import shutil
from datetime import datetime
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import re


def clean_text(text):
    """Clean and normalize text content"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove leading/trailing whitespace
    text = text.strip()
    return text


def extract_text_from_paragraph(para):
    """Extract text from a paragraph, handling runs and formatting"""
    text_parts = []
    for run in para.runs:
        if run.text:
            text_parts.append(run.text)
    
    full_text = ''.join(text_parts)
    return clean_text(full_text)


def paragraph_to_html(para, heading_level=None):
    """Convert a paragraph to HTML"""
    text = extract_text_from_paragraph(para)
    
    if not text:
        return ""
    
    # Check if it's a heading
    if heading_level:
        return f"<h{heading_level}>{text}</h{heading_level}>"
    
    # Check paragraph style for headings
    style_name = para.style.name.lower() if para.style else ''
    if 'heading' in style_name or 'title' in style_name:
        # Extract heading level from style
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            level = min(int(match.group(1)), 6)
            return f"<h{level}>{text}</h{level}>"
        else:
            return f"<h1>{text}</h1>"
    
    # Regular paragraph
    return f"<p>{text}</p>"


def table_to_html(table):
    """Convert a table to HTML"""
    html = "<table>"
    
    for row_idx, row in enumerate(table.rows):
        html += "<tr>"
        for cell in row.cells:
            cell_text = clean_text(cell.text)
            tag = "th" if row_idx == 0 else "td"
            html += f"<{tag}>{cell_text}</{tag}>"
        html += "</tr>"
    
    html += "</table>"
    return html


def extract_images_from_docx(docx_file, output_dir):
    """Extract images from DOCX file and save to output directory"""
    images = []
    doc = Document(docx_file)
    
    # Get image relationships
    relationships = doc.part.rels
    
    for rel_id, rel in relationships.items():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_data = image_part.blob
                image_ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                
                # Generate unique filename
                image_filename = f"image_{len(images)}.{image_ext}"
                image_path = os.path.join(output_dir, image_filename)
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                images.append({
                    'filename': image_filename,
                    'path': image_path,
                    'mimetype': f"image/{image_ext}" if image_ext in ['png', 'jpg', 'jpeg', 'gif'] else "image/png"
                })
            except Exception as e:
                print(f"Warning: Could not extract image {rel_id}: {e}")
    
    return images


def create_epub_from_docx(docx_file, epub_file, include_images=True, preserve_formatting=True, generate_toc=True):
    """
    Create EPUB file from DOCX file
    Returns True if successful, False otherwise
    """
    print(f"Creating EPUB from DOCX: {docx_file}")
    
    try:
        # Read DOCX file
        doc = Document(docx_file)
        
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        title = "Untitled Document"
        try:
            if doc.core_properties.title:
                title = doc.core_properties.title
        except:
            pass
        
        if not title or title == "":
            title = os.path.splitext(os.path.basename(docx_file))[0]
        
        book.set_identifier(os.path.basename(docx_file))
        book.set_title(title)
        
        author = "Unknown Author"
        try:
            if doc.core_properties.author:
                author = doc.core_properties.author
        except:
            pass
        
        book.add_author(author)
        book.set_language('en')
        
        # Create temporary directory for images
        images_dir = os.path.join(os.path.dirname(epub_file), f"images_{os.path.basename(epub_file)}")
        os.makedirs(images_dir, exist_ok=True)
        
        # Extract images if requested
        images = []
        if include_images:
            try:
                images = extract_images_from_docx(docx_file, images_dir)
                print(f"Extracted {len(images)} image(s)")
            except Exception as e:
                print(f"Warning: Could not extract images: {e}")
        
        # Process document content
        chapters = []
        current_chapter_html = []
        chapter_titles = []
        
        # Process all document elements
        for element in doc.element.body:
            # Handle paragraphs
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                para_text = extract_text_from_paragraph(para)
                
                if para_text:
                    style_name = para.style.name.lower() if para.style else ''
                    is_heading = 'heading' in style_name or 'title' in style_name
                    
                    if is_heading and generate_toc:
                        if current_chapter_html:
                            chapters.append('\n'.join(current_chapter_html))
                            current_chapter_html = []
                        
                        chapter_titles.append(para_text)
                        current_chapter_html.append(paragraph_to_html(para))
                    else:
                        current_chapter_html.append(paragraph_to_html(para))
            
            # Handle tables
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                current_chapter_html.append(table_to_html(table))
        
        # Add last chapter
        if current_chapter_html:
            chapters.append('\n'.join(current_chapter_html))
        
        # If no chapters found, create one with all content
        if not chapters:
            print("Warning: No chapters found, creating single chapter")
            all_html = []
            for para in doc.paragraphs:
                para_text = extract_text_from_paragraph(para)
                if para_text:
                    all_html.append(paragraph_to_html(para))
            
            for table in doc.tables:
                all_html.append(table_to_html(table))
            
            chapters.append('\n'.join(all_html))
            chapter_titles = [title] if title else ["Chapter 1"]
        
        print(f"Created {len(chapters)} chapter(s)")
        
        # Add chapters to EPUB
        spine = ['nav']
        toc_items = []
        
        for idx, (chapter_html, chapter_title) in enumerate(zip(chapters, chapter_titles if chapter_titles else [f"Chapter {i+1}" for i in range(len(chapters))])):
            chapter_num = idx + 1
            
            # Create chapter HTML
            html_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{chapter_title}</title>
    <meta charset="utf-8"/>
</head>
<body>
    <h1>{chapter_title}</h1>
    {chapter_html}
</body>
</html>"""
            
            # Create chapter
            chapter = epub.EpubHtml(
                title=chapter_title,
                file_name=f'chapter_{chapter_num}.xhtml',
                lang='en'
            )
            chapter.content = html_content.encode('utf-8')
            book.add_item(chapter)
            spine.append(chapter)
            
            if generate_toc:
                toc_items.append(chapter)
        
        # Add images
        for img in images:
            try:
                with open(img['path'], 'rb') as f:
                    img_item = epub.EpubItem(
                        uid=img['filename'],
                        file_name=f"images/{img['filename']}",
                        media_type=img['mimetype'],
                        content=f.read()
                    )
                    book.add_item(img_item)
            except Exception as e:
                print(f"Warning: Could not add image {img['filename']}: {e}")
        
        # Set table of contents
        if generate_toc and toc_items:
            book.toc = toc_items
        
        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Set spine
        book.spine = spine
        
        # Write EPUB file
        print(f"Writing EPUB file...")
        epub.write_epub(epub_file, book)
        
        # Clean up images directory
        try:
            if os.path.exists(images_dir):
                shutil.rmtree(images_dir)
        except:
            pass
        
        return True
            
    except Exception as e:
        print(f"ERROR: Failed to create EPUB from DOCX: {e}")
        traceback.print_exc()
        return False


def find_ebook_convert():
    """Find ebook-convert binary from Calibre"""
    ebook_convert_paths = [
        'ebook-convert',
        '/usr/bin/ebook-convert',
        '/usr/local/bin/ebook-convert',
        '/opt/calibre/bin/ebook-convert',
        '/opt/calibre/ebook-convert'
    ]
    
    for path in ebook_convert_paths:
        try:
            result = subprocess.run([path, '--version'], capture_output=True, check=True, timeout=5)
            print(f"Found ebook-convert at: {path}")
            return path
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    return None


def convert_epub_to_mobi(epub_file, mobi_file, kindle_optimized=True):
    """
    Convert EPUB file to MOBI format using Calibre ebook-convert
    Returns True if successful, False otherwise
    """
    print(f"Converting EPUB to MOBI: {epub_file} -> {mobi_file}")
    
    # Find ebook-convert
    ebook_convert = find_ebook_convert()
    
    if not ebook_convert:
        print("ERROR: ebook-convert (Calibre) not found. Please ensure Calibre is installed.")
        return False
    
    try:
        # Build conversion command
        convert_cmd = [
            ebook_convert,
            epub_file,
            mobi_file,
            '--language', 'en'
        ]
        
        # MOBI-specific options
        if kindle_optimized:
            convert_cmd.extend([
                '--mobi-file-type', 'both',  # Create both old and new MOBI formats for maximum compatibility
                '--share-not-sync'  # Enable sharing on Kindle
            ])
        else:
            convert_cmd.extend([
                '--mobi-file-type', 'old'  # Old MOBI format only
            ])
        
        print(f"Running ebook-convert: {' '.join(convert_cmd)}")
        
        result = subprocess.run(
            convert_cmd,
            capture_output=True,
            text=True,
            timeout=300  # 5 minute timeout
        )
        
        if result.returncode != 0:
            print(f"ERROR: ebook-convert failed with return code {result.returncode}")
            print(f"stderr: {result.stderr}")
            print(f"stdout: {result.stdout}")
            return False
        
        print("MOBI conversion completed successfully")
        return True
        
    except subprocess.TimeoutExpired:
        print("ERROR: ebook-convert timed out after 5 minutes")
        return False
    except Exception as e:
        print(f"ERROR: Failed to convert EPUB to MOBI: {e}")
        traceback.print_exc()
        return False


def convert_docx_to_mobi(docx_file, output_file, include_images=True, preserve_formatting=True, generate_toc=True, kindle_optimized=True):
    """
    Convert DOCX file to MOBI format
    
    Args:
        docx_file (str): Path to input DOCX file
        output_file (str): Path to output MOBI file
        include_images (bool): Include images from DOCX
        preserve_formatting (bool): Preserve text formatting
        generate_toc (bool): Generate table of contents
        kindle_optimized (bool): Optimize for Kindle devices
        
    Returns:
        bool: True if conversion successful, False otherwise
    """
    print(f"Starting DOCX to MOBI conversion...")
    print(f"Input: {docx_file}")
    print(f"Output: {output_file}")
    print(f"Include images: {include_images}")
    print(f"Preserve formatting: {preserve_formatting}")
    print(f"Generate TOC: {generate_toc}")
    print(f"Kindle optimized: {kindle_optimized}")
    
    try:
        # Check if DOCX file exists
        if not os.path.exists(docx_file):
            print(f"ERROR: DOCX file does not exist: {docx_file}")
            return False
        
        file_size = os.path.getsize(docx_file)
        print(f"DOCX file size: {file_size} bytes")
        
        if file_size == 0:
            print("ERROR: Input file is empty")
            return False
        
        # Create temporary EPUB file
        temp_dir = tempfile.mkdtemp()
        epub_file = os.path.join(temp_dir, os.path.basename(docx_file).replace('.docx', '.epub'))
        
        try:
            # Step 1: Convert DOCX to EPUB
            print("\n=== Step 1: Converting DOCX to EPUB ===")
            if not create_epub_from_docx(docx_file, epub_file, include_images, preserve_formatting, generate_toc):
                print("ERROR: Failed to create EPUB from DOCX")
                return False
            
            if not os.path.exists(epub_file):
                print("ERROR: EPUB file was not created")
                return False
            
            epub_size = os.path.getsize(epub_file)
            print(f"EPUB file created: {epub_size} bytes")
            
            # Step 2: Convert EPUB to MOBI using Calibre
            print("\n=== Step 2: Converting EPUB to MOBI ===")
            if not convert_epub_to_mobi(epub_file, output_file, kindle_optimized):
                print("ERROR: Failed to convert EPUB to MOBI")
                return False
            
            # Verify output file
            if os.path.exists(output_file):
                output_size = os.path.getsize(output_file)
                print(f"\nMOBI file created successfully: {output_size} bytes")
                return True
            else:
                print("ERROR: MOBI file was not created")
                return False
                
        finally:
            # Clean up temporary files
            try:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except:
                pass
            
    except Exception as e:
        print(f"ERROR: Failed to convert DOCX to MOBI: {e}")
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description='Convert DOCX file to MOBI format')
    parser.add_argument('docx_file', help='Path to input DOCX file')
    parser.add_argument('output_file', help='Path to output MOBI file')
    parser.add_argument('--no-images', action='store_true',
                        help='Do not include images from DOCX')
    parser.add_argument('--no-formatting', action='store_true',
                        help='Do not preserve text formatting')
    parser.add_argument('--no-toc', action='store_true',
                        help='Do not generate table of contents')
    parser.add_argument('--no-kindle-optimize', action='store_true',
                        help='Do not optimize for Kindle devices')
    
    args = parser.parse_args()
    
    print("=== DOCX to MOBI Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    success = convert_docx_to_mobi(
        args.docx_file, 
        args.output_file,
        include_images=not args.no_images,
        preserve_formatting=not args.no_formatting,
        generate_toc=not args.no_toc,
        kindle_optimized=not args.no_kindle_optimize
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)


if __name__ == '__main__':
    main()


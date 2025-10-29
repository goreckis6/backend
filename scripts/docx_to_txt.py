#!/usr/bin/env python3
"""
DOCX to TXT Converter
Converts Microsoft Word DOCX files to plain text (TXT) format
Uses Pandoc for clean text output (recommended for TXT)
Falls back to python-docx if Pandoc is not available
"""

import os
import sys
import argparse
import traceback
import subprocess
import tempfile
import shutil
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def find_pandoc():
    """Find Pandoc binary"""
    pandoc_paths = [
        'pandoc',
        '/usr/bin/pandoc',
        '/usr/local/bin/pandoc',
        '/opt/local/bin/pandoc'
    ]
    
    for path in pandoc_paths:
        try:
            result = subprocess.run(
                [path, '--version'],
                capture_output=True,
                check=True,
                timeout=5
            )
            print(f"Found Pandoc at: {path}")
            return path
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    return None


def extract_text_with_docx(docx_file, preserve_line_breaks=True, remove_formatting=True):
    """
    Extract text from DOCX using python-docx (fallback method)
    
    Args:
        docx_file (str): Path to input DOCX file
        preserve_line_breaks (bool): Preserve line breaks
        remove_formatting (bool): Remove formatting (always true for plain text)
        
    Returns:
        str: Extracted text content
    """
    if not HAS_DOCX:
        return None
    
    try:
        print("Using python-docx to extract text from DOCX...")
        doc = Document(docx_file)
        
        text_parts = []
        
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                if preserve_line_breaks:
                    text_parts.append(para_text)
                else:
                    text_parts.append(para_text.replace('\n', ' '))
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(' | '.join(row_texts))
        
        full_text = '\n'.join(text_parts) if preserve_line_breaks else ' '.join(text_parts)
        
        print(f"Extracted {len(full_text)} characters using python-docx")
        return full_text
        
    except Exception as e:
        print(f"Error extracting text with python-docx: {e}")
        traceback.print_exc()
        return None


def convert_docx_to_txt(docx_file, output_file, preserve_line_breaks=True, remove_formatting=True):
    """
    Convert DOCX file to TXT format using Pandoc
    
    Args:
        docx_file (str): Path to input DOCX file
        output_file (str): Path to output TXT file
        preserve_line_breaks (bool): Preserve line breaks from DOCX
        remove_formatting (bool): Remove all formatting (default for clean TXT)
        
    Returns:
        bool: True if conversion successful, False otherwise
    """
    print(f"Starting DOCX to TXT conversion...")
    print(f"Input: {docx_file}")
    print(f"Output: {output_file}")
    print(f"Preserve line breaks: {preserve_line_breaks}")
    print(f"Remove formatting: {remove_formatting}")
    
    try:
        # Check if DOCX file exists
        if not os.path.exists(docx_file):
            print(f"ERROR: DOCX file does not exist: {docx_file}")
            return False
        
        file_size = os.path.getsize(docx_file)
        print(f"DOCX file size: {file_size} bytes")
        
        if file_size == 0:
            print("ERROR: Input file is empty")
            return False
        
        # Find Pandoc
        pandoc = find_pandoc()
        
        if not pandoc:
            print("WARNING: Pandoc not found. Trying fallback method with python-docx...")
            
            # Try fallback method using python-docx
            if HAS_DOCX:
                extracted_text = extract_text_with_docx(docx_file, preserve_line_breaks, remove_formatting)
                
                if extracted_text is not None:
                    # Write extracted text to output file
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(extracted_text)
                    
                    output_size = os.path.getsize(output_file)
                    print(f"TXT file created successfully using python-docx: {output_size} bytes")
                    return True
                else:
                    print("ERROR: Failed to extract text using python-docx")
                    return False
            else:
                print("ERROR: Pandoc not found and python-docx is not available. Please ensure Pandoc is installed or python-docx is in requirements.txt")
                return False
        
        # Create output directory if needed
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Build Pandoc command for clean TXT output
        # Pandoc can convert DOCX to plain text very cleanly
        cmd = [
            pandoc,
            docx_file,
            '-f', 'docx',
            '-t', 'plain',  # Plain text format for clean TXT
            '-o', output_file
        ]
        
        # Add options based on settings
        if preserve_line_breaks:
            # Preserve paragraphs and line breaks
            cmd.extend(['--wrap=none'])  # Don't wrap, preserve original line breaks
        else:
            cmd.extend(['--wrap=preserve'])  # Preserve but may wrap long lines
        
        # Remove formatting is default for plain text, but we can ensure it
        if remove_formatting:
            # Plain text format already removes formatting
            pass
        
        print(f"Running Pandoc: {' '.join(cmd)}")
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=300,  # 5 minute timeout
            encoding='utf-8',
            errors='replace'
        )
        
        if result.stdout:
            print(f"Pandoc stdout: {result.stdout}")
        if result.stderr:
            print(f"Pandoc stderr: {result.stderr}")
        
        # Verify output file exists
        if os.path.exists(output_file):
            output_size = os.path.getsize(output_file)
            print(f"TXT file created successfully: {output_size} bytes")
            return True
        else:
            print(f"ERROR: Pandoc did not create TXT file: {output_file}")
            return False
            
    except subprocess.TimeoutExpired:
        print("ERROR: Pandoc conversion timed out after 5 minutes")
        return False
    except Exception as e:
        print(f"ERROR: Failed to convert DOCX to TXT: {e}")
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description='Convert DOCX file to TXT format using Pandoc')
    parser.add_argument('docx_file', help='Path to input DOCX file')
    parser.add_argument('output_file', help='Path to output TXT file')
    parser.add_argument('--no-line-breaks', action='store_true',
                        help='Do not preserve line breaks')
    parser.add_argument('--keep-formatting', action='store_true',
                        help='Keep formatting (plain text format removes formatting by default)')
    
    args = parser.parse_args()
    
    print("=== DOCX to TXT Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    success = convert_docx_to_txt(
        args.docx_file, 
        args.output_file,
        preserve_line_breaks=not args.no_line_breaks,
        remove_formatting=not args.keep_formatting
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)


if __name__ == '__main__':
    main()


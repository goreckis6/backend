#!/usr/bin/env python3
"""
CSV to DOC Converter
Converts CSV files to Microsoft Word DOC format using python-docx library.
This is a fallback solution when LibreOffice fails due to Java dependency issues.
"""

import pandas as pd
import argparse
import os
import sys
from datetime import datetime
import traceback

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
except ImportError as e:
    print(f"ERROR: Required DOCX library not available: {e}")
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

def create_doc_from_csv(csv_file, output_file, title="CSV Data", author="Unknown"):
    """
    Convert CSV file to DOC format using python-docx.
    
    Args:
        csv_file (str): Path to input CSV file
        output_file (str): Path to output DOC file
        title (str): Document title
        author (str): Document author
    """
    print(f"Starting CSV to DOC conversion...")
    print(f"Input: {csv_file}")
    print(f"Output: {output_file}")
    print(f"Title: {title}")
    print(f"Author: {author}")
    
    try:
        # Read CSV file
        print("Reading CSV file...")
        df = pd.read_csv(csv_file)
        print(f"CSV loaded successfully: {len(df)} rows, {len(df.columns)} columns")
        
        # Create DOCX document
        print("Creating DOCX document...")
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()
        
        # Add title
        print("Adding document title...")
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add author and date info
        author_para = doc.add_paragraph(f"Author: {author}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add empty line
        doc.add_paragraph()
        
        # Create table
        print("Creating data table...")
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the table
        table.style = 'Table Grid'
        
        # Add header row
        print("Adding table headers...")
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            # Style header cells
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        print(f"Adding {len(df)} data rows...")
        for idx, row in df.iterrows():
            if idx % 100 == 0:
                print(f"Processing row {idx + 1}/{len(df)}")
            
            # Add new row
            table.add_row()
            row_cells = table.rows[-1].cells
            
            for i, value in enumerate(row):
                # Handle NaN values and convert to string
                cell_value = str(value) if pd.notna(value) else ""
                row_cells[i].text = cell_value
                
                # Style data cells
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add summary
        print("Adding document summary...")
        doc.add_paragraph()
        summary_para = doc.add_paragraph(f"Summary: {len(df)} rows, {len(df.columns)} columns")
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        print(f"Saving DOCX document to {output_file}...")
        doc.save(output_file)
        
        # Verify file was created
        if os.path.exists(output_file):
            file_size = os.path.getsize(output_file)
            print(f"DOCX file created successfully: {file_size} bytes")
            return True
        else:
            print("ERROR: DOCX file was not created")
            return False
            
    except Exception as e:
        print(f"ERROR: Failed to create DOCX from CSV: {e}")
        traceback.print_exc()
        return False

def main():
    parser = argparse.ArgumentParser(description='Convert CSV to DOC format')
    parser.add_argument('csv_file', help='Input CSV file path')
    parser.add_argument('output_file', help='Output DOC file path')
    parser.add_argument('--title', default='CSV Data', help='Document title')
    parser.add_argument('--author', default='Unknown', help='Document author')
    
    args = parser.parse_args()
    
    print("=== CSV to DOC Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    # Check if input file exists
    if not os.path.exists(args.csv_file):
        print(f"ERROR: Input CSV file not found: {args.csv_file}")
        sys.exit(1)
    
    # Check pandas availability
    try:
        print(f"Pandas version: {pd.__version__}")
    except Exception as e:
        print(f"ERROR: Pandas not available: {e}")
        sys.exit(1)
    
    # Check python-docx availability
    try:
        from docx import Document
        print("python-docx library is available")
    except Exception as e:
        print(f"ERROR: python-docx not available: {e}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(args.output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Convert CSV to DOCX
    success = create_doc_from_csv(
        args.csv_file,
        args.output_file,
        args.title,
        args.author
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)

if __name__ == "__main__":
    main()

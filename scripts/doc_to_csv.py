#!/usr/bin/env python3
"""
DOC to CSV Converter
Converts Microsoft Word DOC files to CSV format
Uses LibreOffice for best table-based DOC conversion
"""

import os
import sys
import argparse
import traceback
import subprocess
import tempfile
import shutil
import csv
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def find_libreoffice():
    """Find LibreOffice binary"""
    libreoffice_paths = [
        'libreoffice',
        '/usr/bin/libreoffice',
        '/usr/local/bin/libreoffice',
        '/opt/libreoffice/program/soffice',
        'soffice'
    ]
    
    for path in libreoffice_paths:
        try:
            result = subprocess.run(
                [path, '--version'],
                capture_output=True,
                check=True,
                timeout=5
            )
            print(f"Found LibreOffice at: {path}")
            return path
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    return None


def convert_doc_to_csv_libreoffice(doc_file, output_file, delimiter=',', extract_tables=True, include_paragraphs=True):
    """
    Convert DOC file to CSV format using LibreOffice
    Strategy: DOC -> CSV via LibreOffice (best for table-based DOC files)
    
    Args:
        doc_file (str): Path to input DOC file
        output_file (str): Path to output CSV file
        delimiter (str): CSV delimiter (comma, semicolon, tab, pipe)
        extract_tables (bool): Extract tables from DOC
        include_paragraphs (bool): Include paragraphs as data rows
    
    Returns:
        bool: True if conversion successful, False otherwise
    """
    print(f"Starting DOC to CSV conversion using LibreOffice...")
    print(f"Input: {doc_file}")
    print(f"Output: {output_file}")
    print(f"Delimiter: {delimiter}")
    print(f"Extract tables: {extract_tables}")
    print(f"Include paragraphs: {include_paragraphs}")
    
    try:
        # Check if DOC file exists
        if not os.path.exists(doc_file):
            print(f"ERROR: DOC file does not exist: {doc_file}")
            return False
        
        file_size = os.path.getsize(doc_file)
        print(f"DOC file size: {file_size} bytes")
        
        if file_size == 0:
            print("ERROR: Input file is empty")
            return False
        
        # Find LibreOffice
        libreoffice = find_libreoffice()
        
        if not libreoffice:
            print("ERROR: LibreOffice not found. Please ensure LibreOffice is installed.")
            return False
        
        # Create temporary directory for conversion
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Create output directory if needed
            output_dir = os.path.dirname(output_file)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Strategy: DOC -> DOCX (using LibreOffice) -> CSV (using python-docx and pandas)
            # Step 1: Convert DOC to DOCX using LibreOffice
            base_name = os.path.splitext(os.path.basename(doc_file))[0]
            intermediate_docx = os.path.join(temp_dir, f"{base_name}.docx")
            
            print("Step 1: Converting DOC to DOCX using LibreOffice...")
        cmd = [
                libreoffice,
            '--headless',
            '--invisible',
            '--nocrashreport',
            '--nodefault',
            '--nofirststartwizard',
            '--nolockcheck',
            '--nologo',
            '--norestore',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
            doc_file
        ]
        
            # Set LibreOffice environment
        env = os.environ.copy()
        env['SAL_USE_VCLPLUGIN'] = 'svp'
        env['HOME'] = '/tmp'
        env['LANG'] = 'en_US.UTF-8'
        env['LC_ALL'] = 'en_US.UTF-8'
        
            print(f"Running LibreOffice: {' '.join(cmd)}")
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
                timeout=300,  # 5 minute timeout
            env=env,
            encoding='utf-8',
            errors='replace'
        )
        
        if result.stdout:
                print(f"LibreOffice stdout: {result.stdout}")
            if result.stderr and 'Error:' not in result.stderr:
                print(f"LibreOffice stderr: {result.stderr}")
            
            # Check if DOCX was created
            if not os.path.exists(intermediate_docx):
                print(f"ERROR: LibreOffice did not create DOCX file: {intermediate_docx}")
                print(f"Contents of temp directory {temp_dir}: {os.listdir(temp_dir)}")
            return False
            
            print(f"Step 1 complete: DOCX created at {intermediate_docx}")
            
            # Step 2: Extract tables from DOCX and convert to CSV
            if not HAS_DOCX:
                print("ERROR: python-docx is required but not available. Please install python-docx.")
                return False
            
            print("Step 2: Extracting tables from DOCX...")
            doc = Document(intermediate_docx)
            
            # Collect all table data
            all_rows = []
            
            if extract_tables and len(doc.tables) > 0:
                print(f"Found {len(doc.tables)} table(s) in document")
                for table_idx, table in enumerate(doc.tables):
                    print(f"Processing table {table_idx + 1}...")
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
                            row_data.append(cell_text)
                        if any(cell for cell in row_data):  # Only add non-empty rows
                            all_rows.append(row_data)
            
            # Include paragraphs if requested and no tables found
            if include_paragraphs and len(all_rows) == 0:
                print("No tables found, extracting paragraphs...")
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        all_rows.append([para_text])
            
            if len(all_rows) == 0:
                print("WARNING: No data extracted from document")
                # Create empty CSV file
                with open(output_file, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f, delimiter=delimiter)
                    pass  # Empty file
            else:
                # Write to CSV with specified delimiter
                print(f"Writing {len(all_rows)} rows to CSV...")
                with open(output_file, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f, delimiter=delimiter)
                    for row in all_rows:
                        writer.writerow(row)
        
        # Verify output file
            if os.path.exists(output_file):
                output_size = os.path.getsize(output_file)
                print(f"CSV file created successfully: {output_size} bytes, {len(all_rows)} rows")
                return True
            else:
                print(f"ERROR: CSV file was not created at {output_file}")
                return False
                
        finally:
            # Clean up temp directory
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"Warning: Could not clean up temp directory: {e}")
            
    except subprocess.TimeoutExpired:
        print("ERROR: LibreOffice conversion timed out after 5 minutes")
        return False
    except Exception as e:
        print(f"ERROR: Failed to convert DOC to CSV: {e}")
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description='Convert DOC file to CSV format using LibreOffice')
    parser.add_argument('doc_file', help='Path to input DOC file')
    parser.add_argument('output_file', help='Path to output CSV file')
    parser.add_argument('--delimiter', default=',', choices=[',', ';', '\t', '|'],
                        help='CSV delimiter (default: comma)')
    parser.add_argument('--no-tables', action='store_true',
                        help='Do not extract tables (LibreOffice extracts tables by default)')
    parser.add_argument('--no-paragraphs', action='store_true',
                        help='Do not include paragraphs (not applicable for LibreOffice CSV export)')
    
    args = parser.parse_args()
    
    print("=== DOC to CSV Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    success = convert_doc_to_csv_libreoffice(
        args.doc_file, 
        args.output_file,
        delimiter=args.delimiter,
        extract_tables=not args.no_tables,
        include_paragraphs=not args.no_paragraphs
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)


if __name__ == '__main__':
    main()

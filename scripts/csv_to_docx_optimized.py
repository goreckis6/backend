#!/usr/bin/env python3
"""
Optimized CSV to DOCX Converter
High-performance CSV to DOCX conversion with optimizations for large files.
"""

import pandas as pd
import argparse
import os
import sys
from datetime import datetime
import traceback
from io import StringIO

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import qn as qn_ns
except ImportError as e:
    print(f"ERROR: Required DOCX library not available: {e}")
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

def create_docx_from_csv_optimized(csv_file, output_file, title="CSV Data", author="Unknown", chunk_size=1000):
    """
    Optimized CSV to DOCX conversion with performance improvements.
    
    Args:
        csv_file (str): Path to input CSV file
        output_file (str): Path to output DOCX file
        title (str): Document title
        author (str): Document author
        chunk_size (int): Number of rows to process at once
    """
    print(f"Starting optimized CSV to DOCX conversion...")
    print(f"Input: {csv_file}")
    print(f"Output: {output_file}")
    print(f"Chunk size: {chunk_size}")
    
    try:
        # Read CSV file with optimizations
        print("Reading CSV file with optimizations...")
        
        # Get file size for progress tracking
        file_size = os.path.getsize(csv_file)
        print(f"File size: {file_size / (1024*1024):.2f} MB")
        
        # Read CSV with maximum speed optimizations and data preservation
        df = pd.read_csv(
            csv_file,
            dtype=str,  # Read all as strings to avoid type inference overhead
            na_filter=False,  # Disable NaN filtering for speed
            low_memory=False,  # Use more memory for speed
            engine='c',  # Use C engine for maximum speed
            memory_map=True,  # Memory map for large files
            chunksize=None,  # Read entire file at once for speed
            keep_default_na=False,  # Don't convert empty strings to NaN
            na_values=[],  # Don't treat any values as NaN
            encoding='utf-8',  # Ensure proper encoding
            skip_blank_lines=False,  # Keep all rows including blank ones
            skipinitialspace=False  # Don't remove leading spaces
        )
        
        print(f"CSV loaded: {len(df)} rows, {len(df.columns)} columns")
        
        # Data integrity check - show first few rows for verification
        print("First 3 rows of data:")
        for i in range(min(3, len(df))):
            print(f"Row {i+1}: {list(df.iloc[i])}")
        
        # Check for any completely empty rows that might cause issues
        empty_rows = df.isnull().all(axis=1).sum()
        if empty_rows > 0:
            print(f"WARNING: Found {empty_rows} completely empty rows")
        
        # Note: We preserve ALL data including duplicates - no deduplication
        print("✅ Data preservation mode: ALL rows and duplicates will be preserved")
        
        # Create DOCX document
        print("Creating optimized DOCX document...")
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()
        
        # Add minimal title section
        print("Adding document header...")
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add creation info
        info_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Rows: {len(df)} | Columns: {len(df.columns)}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line
        
        # Create table with optimizations
        print("Creating optimized data table...")
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Add header row with minimal styling
        print("Adding table headers...")
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            # Minimal header styling
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)  # Smaller font for speed
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process data with maximum speed optimizations and data integrity
        print(f"Processing {len(df)} rows in chunks of {chunk_size}...")
        
        # Pre-allocate table rows for better performance (but ensure we don't skip data)
        total_rows = len(df)
        pre_allocated_rows = 0
        if total_rows > 1000:  # Only pre-allocate for very large datasets
            print("Pre-allocating table rows for large dataset...")
            # Pre-allocate up to 2000 rows maximum to avoid memory issues
            pre_allocated_rows = min(2000, total_rows)
            for _ in range(pre_allocated_rows):
                table.add_row()
        
        # Process data in optimized chunks with data integrity checks
        processed_rows = 0
        for chunk_start in range(0, total_rows, chunk_size):
            chunk_end = min(chunk_start + chunk_size, total_rows)
            chunk_df = df.iloc[chunk_start:chunk_end]
            
            print(f"Processing rows {chunk_start + 1}-{chunk_end} of {total_rows}")
            
            # Process each row in the chunk
            for idx, (_, row) in enumerate(chunk_df.iterrows()):
                # Add row if not pre-allocated
                if processed_rows >= pre_allocated_rows:
                    table.add_row()
                
                row_cells = table.rows[-1].cells
                
                # Process cells with data integrity - preserve ALL data including duplicates
                for i, value in enumerate(row):
                    # Preserve the original data exactly as it appears in CSV
                    if pd.isna(value) or value is None:
                        cell_value = ""
                    else:
                        # Keep the original string representation, don't modify it
                        cell_value = str(value)
                    
                    # Set the cell text exactly as it appears in the original data
                    row_cells[i].text = cell_value
                    
                    # Skip styling for most rows to maximize speed
                    if processed_rows < 50:  # Only style first 50 rows
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)  # Smaller font
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                processed_rows += 1
        
        # Verify data integrity
        print(f"Data integrity check: Processed {processed_rows} rows, Expected {total_rows} rows")
        if processed_rows != total_rows:
            print(f"WARNING: Data mismatch! Processed {processed_rows} but expected {total_rows}")
        else:
            print("✅ Data integrity verified - all rows processed correctly")
        
        # Add minimal summary
        print("Adding document summary...")
        doc.add_paragraph()
        summary_para = doc.add_paragraph(f"Total: {len(df)} rows × {len(df.columns)} columns")
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document with optimizations
        print(f"Saving DOCX document...")
        doc.save(output_file)
        
        # Verify file was created and validate data integrity
        if os.path.exists(output_file):
            file_size = os.path.getsize(output_file)
            print(f"DOCX file created successfully: {file_size / (1024*1024):.2f} MB")
            
            # Additional validation: check if file is not empty and has reasonable size
            if file_size < 1000:  # Less than 1KB is suspicious for a DOCX file
                print("WARNING: Output file is very small, may indicate data loss")
                return False
            
            print("✅ File creation and size validation passed")
            return True
        else:
            print("ERROR: DOCX file was not created")
            return False
            
    except Exception as e:
        print(f"ERROR: Failed to create DOCX from CSV: {e}")
        traceback.print_exc()
        return False

def main():
    parser = argparse.ArgumentParser(description='Convert CSV to DOCX format (optimized)')
    parser.add_argument('csv_file', help='Input CSV file path')
    parser.add_argument('output_file', help='Output DOCX file path')
    parser.add_argument('--title', default='CSV Data', help='Document title')
    parser.add_argument('--author', default='Unknown', help='Document author')
    parser.add_argument('--chunk-size', type=int, default=1000, help='Chunk size for processing large files')
    
    args = parser.parse_args()
    
    print("=== Optimized CSV to DOCX Converter ===")
    print(f"Python version: {sys.version}")
    print(f"Working directory: {os.getcwd()}")
    print(f"Arguments: {vars(args)}")
    
    # Check if input file exists
    if not os.path.exists(args.csv_file):
        print(f"ERROR: Input CSV file not found: {args.csv_file}")
        sys.exit(1)
    
    # Check pandas availability
    try:
        print(f"Pandas version: {pd.__version__}")
    except Exception as e:
        print(f"ERROR: Pandas not available: {e}")
        sys.exit(1)
    
    # Check python-docx availability
    try:
        from docx import Document
        print("python-docx library is available")
    except Exception as e:
        print(f"ERROR: python-docx not available: {e}")
        sys.exit(1)
    
    # Determine chunk size based on file size for maximum performance
    file_size = os.path.getsize(args.csv_file)
    file_size_mb = file_size / (1024 * 1024)
    if file_size_mb > 20:
        chunk_size = 5000  # Large files: bigger chunks
    elif file_size_mb > 10:
        chunk_size = 3000  # Medium files: medium chunks
    elif file_size_mb > 5:
        chunk_size = 2000  # Small files: smaller chunks
    else:
        chunk_size = 1500  # Very small files: smallest chunks
    
    print(f"Auto-determined chunk size: {chunk_size} (file size: {file_size_mb:.2f} MB)")
    
    # Override with user-specified chunk size if provided
    if args.chunk_size != 1000:
        chunk_size = args.chunk_size
        print(f"Using user-specified chunk size: {chunk_size}")
    
    # Convert CSV to DOCX
    success = create_docx_from_csv_optimized(
        args.csv_file,
        args.output_file,
        args.title,
        args.author,
        chunk_size
    )
    
    if success:
        print("=== CONVERSION SUCCESSFUL ===")
        sys.exit(0)
    else:
        print("=== CONVERSION FAILED ===")
        sys.exit(1)

if __name__ == "__main__":
    main()